"""One-off builder for instructions for editors.docx."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Base style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def h1(text):
    p = doc.add_heading(text, level=1)
    return p


def h2(text):
    p = doc.add_heading(text, level=2)
    return p


def h3(text):
    p = doc.add_heading(text, level=3)
    return p


def p(text, bold=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    if bold:
        run.bold = True
    return para


def bullet(text):
    para = doc.add_paragraph(style="List Bullet")
    para.add_run(text)
    return para


def numbered(text):
    para = doc.add_paragraph(style="List Number")
    para.add_run(text)
    return para


def code_block(text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return para


def inline_code(para, text):
    run = para.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    return run


# =============================================================================
# TITLE
# =============================================================================
title = doc.add_heading("Instructions for Editors — Wide Tim Site", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT

p(
    "This document explains how to replace placeholder images and update the "
    "text content of the Wide Tim site. You do not need to be a programmer to "
    "do this — but you do need to be careful with punctuation. Always preview "
    "your changes locally (see the “Previewing Your Changes” section) before "
    "publishing."
)

# =============================================================================
# BEFORE YOU START
# =============================================================================
h1("1. Before You Start")

bullet("The site lives in: /Users/margaretzheng/Desktop/wide_tim")
bullet(
    "You need Node.js installed on your computer. If you do not have it, "
    "download it from https://nodejs.org (use the LTS version)."
)
bullet(
    "Open the Terminal app, navigate to the project folder, and run "
    "“npm install” once. This downloads everything the site needs to run."
)

para = doc.add_paragraph()
para.add_run("In Terminal:").bold = True
code_block(
    "cd /Users/margaretzheng/Desktop/wide_tim\n"
    "npm install"
)

p(
    "You only need to run “npm install” once (or again if dependencies change). "
    "After that, all your day-to-day work is just editing files and previewing."
)

# =============================================================================
# HOW THE SITE IS ORGANIZED
# =============================================================================
h1("2. How the Site Is Organized")

p(
    "You will only need to touch two areas of the project: the content files "
    "(where the text lives) and the image folders. Here is a quick map:"
)

bullet(
    "src/content/ — Text content for each page. ONE FILE PER PAGE. "
    "This is where 90% of your edits will happen."
)
bullet(
    "public/ — Images. Each page has its own image folder, "
    "e.g. /public/home-page-images/, /public/about-page-images/."
)
bullet(
    "images-src/ — Source images for the HOME PAGE cards only. "
    "These are auto-cropped into squares by a build script."
)
bullet(
    "src/pages/ — Page templates (the HTML structure). "
    "Do not edit these unless you really know what you are doing."
)

h2("Which content file matches which page?")

mapping = [
    ("home.ts", "Home page (the landing page with the card grid)"),
    ("aboutTim.ts", "“About Wide Tim” page"),
    ("aboutCreator.ts", "“About the Creator” page"),
    ("projects.ts", "“Selected Projects” page"),
    ("collaborators.ts + collaborators.csv", "“Collaborations” page"),
    ("press.ts", "“Press” page"),
    ("fun.ts", "“Just Here for Fun” page"),
    ("contact.ts", "Contact page"),
    ("socials.ts", "Social media links (footer)"),
    ("beingWide.ts", "“Being Wide 2024” sub-page"),
    ("carousel.ts", "“Make a Wider Tim” carousel content"),
    ("site.ts", "Site-wide settings (page title, navigation)"),
]
for fname, desc in mapping:
    para = doc.add_paragraph(style="List Bullet")
    inline_code(para, fname)
    para.add_run(" — " + desc)

# =============================================================================
# EDITING TEXT CONTENT
# =============================================================================
h1("3. Editing Text Content")

p(
    "Open any file in src/content/ with a plain text editor (TextEdit in plain "
    "text mode is fine; VS Code is even better). Each file is a TypeScript "
    "object — essentially a labeled list of fields. You will see things like:"
)

code_block(
    'export const aboutCreator = {\n'
    '  heading: "About the Creator",\n'
    '  name: "Tianyuan \\u201CMargaret\\u201D Zheng",\n'
    '  facts: [\n'
    '    "MIT \'23, MEng \'24",\n'
    '    "Course 18, 6-14",\n'
    '  ],\n'
    '  paragraphs: [\n'
    '    `Margaret is the artist behind Wide Tim...`,\n'
    '  ],\n'
    '};\n'
)

h2("The three kinds of text you will edit")

para = doc.add_paragraph(style="List Number")
para.add_run("Double-quoted strings — e.g. ").bold = False
inline_code(para, '"About the Creator"')
para.add_run(
    ". Edit the text BETWEEN the quotes. Plain text only — no HTML."
)

para = doc.add_paragraph(style="List Number")
para.add_run("Backtick strings — e.g. ")
inline_code(para, "`Margaret is the artist...`")
para.add_run(
    ". These ALLOW basic HTML, so you can put links and bold inline. "
    'Example: `She started <a href="https://...">here</a>.`'
)

para = doc.add_paragraph(style="List Number")
para.add_run("Arrays (lists in square brackets ")
inline_code(para, "[ ... ]")
para.add_run(
    "). Each item is separated by a comma. To add a new item, copy an "
    "existing item, paste it on a new line, and edit it. To remove an item, "
    "delete the whole line including its trailing comma."
)

h2("Rules to follow")

bullet(
    "Never delete the surrounding braces { } or the brackets [ ] — they "
    "define the shape of the data."
)
bullet(
    "Always keep the trailing comma at the end of each line inside { } or [ ]. "
    "Missing commas are the #1 cause of build errors."
)
bullet(
    "Inside backtick strings, if you use HTML, ALWAYS close your tags: "
    '<a href="...">text</a>, not <a href="...">text.'
)
bullet(
    "If your text contains a real apostrophe (e.g. “Margaret’s”) inside a "
    "double-quoted string, use a curly apostrophe (’) OR escape it: "
    '"Margaret\\\'s". Backtick strings handle apostrophes without trouble.'
)

# =============================================================================
# REPLACING PLACEHOLDER IMAGES
# =============================================================================
h1("4. Replacing Placeholder Images")

p(
    "There are TWO different procedures depending on which page you are "
    "working on. The home page is special because its card thumbnails are "
    "auto-cropped to squares. Every other page uses images exactly as you "
    "save them."
)

h2("4a. Home page card thumbnails (auto-processed)")

para = doc.add_paragraph(style="List Number")
para.add_run("Place the RAW image (any size, any aspect ratio) into the ")
inline_code(para, "/images-src/")
para.add_run(" folder.")

para = doc.add_paragraph(style="List Number")
para.add_run("Name it to match the filename already referenced in ")
inline_code(para, "src/content/home.ts")
para.add_run(
    ". For example, the “About Wide Tim” card currently points at "
)
inline_code(para, '/home-page-images/about-wide-tim.png')
para.add_run(", so name your replacement ")
inline_code(para, "about-wide-tim.png")
para.add_run(" and drop it into ")
inline_code(para, "/images-src/")
para.add_run(".")

para = doc.add_paragraph(style="List Number")
para.add_run("Run ")
inline_code(para, "npm run dev")
para.add_run(" (or ")
inline_code(para, "npm run images")
para.add_run(
    "). The script crops the image to a 1200×1200 square (centered) and "
    "writes the result into "
)
inline_code(para, "/public/home-page-images/")
para.add_run(". The home page automatically picks it up.")

p(
    "If you want to use a NEW filename instead, change the “image:” line in "
    "home.ts to match, AND put the source image in /images-src/ with the same "
    "name."
)

h2("4b. All other pages (manual)")

para = doc.add_paragraph(style="List Number")
para.add_run("Find the page’s image folder under ")
inline_code(para, "/public/")
para.add_run(". Examples:")

bullet("About the Creator → /public/about-creator-images/  (create the folder if it does not exist yet)")
bullet("About Wide Tim    → /public/about-page-images/")
bullet("Selected Projects → /public/projects-images/  (create if needed)")
bullet("Just Here for Fun → /public/fun-images/  (create if needed)")
p(
    "Each page’s content file (src/content/*.ts) has a comment near the "
    "“image:” field telling you exactly which folder it expects."
)

para = doc.add_paragraph(style="List Number")
para.add_run(
    "Drop your finished, cropped image into the folder. Use the final "
    "dimensions and aspect ratio you want displayed — unlike the home page, "
    "these images are shown as-is."
)

para = doc.add_paragraph(style="List Number")
para.add_run("Open the matching content file in ")
inline_code(para, "src/content/")
para.add_run(" and set the “image:” field to the path to your image. Example, in ")
inline_code(para, "aboutCreator.ts")
para.add_run(":")

code_block('image: "/about-creator-images/portrait.jpg",')

para = doc.add_paragraph(style="List Number")
para.add_run("Also fill in ")
inline_code(para, "imageAlt:")
para.add_run(
    " with a short description of what is in the image. This is read aloud "
    "by screen readers and shown if the image fails to load."
)
code_block('imageAlt: "Margaret holding a Wide Tim plush",')

h2("4c. Image file tips")

bullet("Use lowercase filenames with hyphens — no spaces, no capital letters. Good: class-of-2028-poster.jpg. Bad: Class Of 2028 Poster.JPG.")
bullet("Keep file sizes small (ideally under ~500 KB). Large images make the site slow.")
bullet("Use .jpg for photos, .png for illustrations with sharp edges or transparency.")
bullet("Home-page sources can be any aspect ratio — they get center-cropped to square automatically.")
bullet("All other pages display images as-is, so crop them yourself before saving.")

# =============================================================================
# ADDING NEW CONTENT
# =============================================================================
h1("5. Adding New Content")

h2("5a. Add a new project (src/content/projects.ts)")

p(
    "Copy an existing project block and paste a new one inside the “projects” "
    "array. Newest first is the convention. Only the “name” and “date” fields "
    "are required; everything else is optional."
)
code_block(
    "{\n"
    '  name: "Pie with Wide Tim",\n'
    '  date: "3.14.2025",\n'
    '  url: "https://artfinity.mit.edu/event/bright-spot-pie-with-wide-tim",\n'
    '  description: "Short blurb about the project.",\n'
    '  image: "/projects-images/pie.jpg",\n'
    '  imageAlt: "Wide Tim with a pie",\n'
    "},"
)
p("Do not forget the comma after the closing brace.")

h2("5b. Add a new collaborator / event (src/content/collaborators.csv)")

p(
    "Open collaborators.csv in any spreadsheet app (Numbers, Excel, Google "
    "Sheets) OR in a plain-text editor. Add a row with these columns:"
)
bullet("collaborator — the club or office name")
bullet("event — the event name (leave blank if not tied to an event)")
bullet("month — number 1–12")
bullet("year — four-digit year")
bullet("day — day of the month")
bullet("group — optional grouping label")
bullet("url — optional link")
bullet("note — optional free-text note")
p("Save as CSV (not .xlsx). The site will pick up your new row automatically.")

h2("5c. Add a new quote, bullet, or paragraph")

p(
    "Open the relevant content file, find the array (named things like "
    "“quotes”, “inPersonBullets”, “paragraphs”), and add a new entry that "
    "matches the shape of the existing ones. Copy + paste + edit is the "
    "safest way to do this."
)

h2("5d. Add a new social link or contact line")

p(
    "Edit src/content/socials.ts (for the footer links) or src/content/"
    "contact.ts (for the Contact page). Both files contain arrays of "
    "objects with “label”, “handle”, and “url” fields."
)

# =============================================================================
# PREVIEWING
# =============================================================================
h1("6. Previewing Your Changes")

para = doc.add_paragraph(style="List Number")
para.add_run("Open Terminal and navigate to the project folder:")
code_block("cd /Users/margaretzheng/Desktop/wide_tim")

para = doc.add_paragraph(style="List Number")
para.add_run("Start the dev server:")
code_block("npm run dev")

para = doc.add_paragraph(style="List Number")
para.add_run(
    "It will print a URL — usually http://localhost:4321. Open it in your browser."
)

para = doc.add_paragraph(style="List Number")
para.add_run(
    "Every time you save a content file or add an image, the page reloads "
    "automatically. No need to restart the server for content changes."
)

para = doc.add_paragraph(style="List Number")
para.add_run("Press Ctrl+C in the Terminal to stop the server when you are done.")

# =============================================================================
# COMMON PITFALLS
# =============================================================================
h1("7. Common Pitfalls & Troubleshooting")

bullet(
    "Red error in the terminal after saving a content file → almost always a "
    "missing comma, missing quote, or unbalanced bracket. Compare your edited "
    "block against a neighboring block."
)
bullet(
    "Image does not show up → double-check that the path starts with “/” and "
    "matches the folder name and filename EXACTLY (it is case-sensitive)."
)
bullet(
    "Home-page card still shows the placeholder → confirm the source image is "
    "in /images-src/ (not somewhere else), then re-run npm run dev. The "
    "processor only runs at startup or when you run “npm run images”."
)
bullet(
    "A link inside a backtick string is broken → make sure you closed the "
    'tag: <a href="https://...">text</a>. The closing </a> matters.'
)
bullet(
    "Changes do not appear in the browser → make sure you saved the file, "
    "and that npm run dev is still running. If in doubt, refresh the browser."
)

# =============================================================================
# PUBLISHING
# =============================================================================
h1("8. Publishing the Site")

p(
    "When your changes look good locally, you build a “production” copy of "
    "the site and upload it to the web host."
)
para = doc.add_paragraph(style="List Number")
para.add_run("Build the site:")
code_block("npm run build")
p("This creates a /dist/ folder containing the finished static site.")

para = doc.add_paragraph(style="List Number")
para.add_run(
    "Upload the contents of /dist/ to the web host. (Ask Margaret which host "
    "and how to upload — this depends on where the site is currently deployed.)"
)

p(
    "That is everything. When in doubt: preview locally, take small steps, "
    "and keep an eye on the Terminal for red error messages."
)

# =============================================================================
# QUICK REFERENCE CARD
# =============================================================================
h1("Appendix: Quick Reference")

para = doc.add_paragraph()
para.add_run("Replace a home-page card thumbnail: ").bold = True
para.add_run("drop image into /images-src/ with the same filename → npm run dev.")

para = doc.add_paragraph()
para.add_run("Replace any other image: ").bold = True
para.add_run(
    "drop image into the page’s /public/<page>-images/ folder → update the "
    "image: field in that page’s src/content/*.ts file."
)

para = doc.add_paragraph()
para.add_run("Edit page text: ").bold = True
para.add_run(
    "open src/content/<page>.ts, edit text between quotes or backticks, save."
)

para = doc.add_paragraph()
para.add_run("Preview: ").bold = True
inline_code(para, "npm run dev")
para.add_run(" — open http://localhost:4321.")

para = doc.add_paragraph()
para.add_run("Publish: ").bold = True
inline_code(para, "npm run build")
para.add_run(" — upload /dist/ to the web host.")

doc.save("/Users/margaretzheng/Desktop/wide_tim/instructions/instructions for editors.docx")
print("OK")
